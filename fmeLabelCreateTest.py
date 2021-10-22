import fme, fmeobjects, os, docx, re
from docx.shared import Pt, Inches
from docx.table import _Cell
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.shared import OxmlElement, qn


# Helper function to set cell margins via https://stackoverflow.com/questions/51060431/how-to-set-cell-margins-of-tables-in-ms-word-using-python-docx
def set_cell_margins(cell: _Cell, **kwargs):
    """
    cell:  actual cell instance you want to modify
    usage:
        set_cell_margins(cell, top=50, start=50, bottom=50, end=50)
    provided values are in twentieths of a point (1/1440 of an inch).
    read more here: http://officeopenxml.com/WPtableCellMargins.php
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for m in [
        "top",
        "start",
        "bottom",
        "end",
    ]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

    tcPr.append(tcMar)


class LabelCreate(object):
    def __init__(self):
        self.featureList = []

    def input(self, feature):
        self.featureList.append(feature)

    def close(self):
        # Populate dataset with empty strings until we have enough to fill in to the end of the last page/table with the correct cell sizes
        while len(self.featureList) % 14:
            blankFeature = fmeobjects.FMEFeature()
            blankFeature.setGeometry(None)
            blankFeature.setAttribute('ContainerComponentA', '')
            blankFeature.setAttribute('ContainerComponentB', '')
            blankFeature.setAttribute('PortNumberA', '')
            blankFeature.setAttribute('PortNumberB', '')
            blankFeature.setAttribute('RackNameA', '')
            blankFeature.setAttribute('RackNameB', '')
            blankFeature.setAttribute('SlotPortA', '')
            blankFeature.setAttribute('SlotPortB', '')
            self.featureList.append(blankFeature)

        # Set directory to save the document
        os.chdir(os.getcwd())

        # Create a document object
        document = docx.Document()

        # Set some variables to keep track of how many cells we have written out on and current row/column index
        cellsNum = 0
        rowNum = 0
        colNum = 0

        # Build out the table by iterating over the data
        for i in self.featureList:
            # Add a new table to the document if this is the first data item, or every 52 cells
            if not cellsNum % 51:
                # For tables beyond the first, we also need to add a page break
                if cellsNum >= 51:
                    document.add_page_break()
                # Add table, set styles, reset row index and cell index for starting a new table (colNum reset at end of last table)
                table = document.add_table(4, 13, document.styles['Table Grid'])
                table.style.font.name = 'Arial'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False
                cellsNum = 0
                rowNum = 0
            # Rows with an even index will get filled with data and set to 1 inch wide
            if not rowNum % 2:
                table.rows[rowNum].height = Inches(1)
                table.rows[rowNum].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                row = table.rows[rowNum]
                # Access the current cell, write the label info, set the width, and increment the current cell number by 1
                cell = table.cell(rowNum, colNum)
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run(
                    f'{i.getAttribute("ContainerComponentA")}\n{i.getAttribute("RackNameA")}\n{i.getAttribute("SlotPortA")}\n{i.getAttribute("ContainerComponentB")}\n{i.getAttribute("RackNameB")}\n{i.getAttribute("SlotPortB")}')
                # Set font size to 6.5pt if strings will take up multiple lines more than 4 times, 7pt if more than 2 times, otherwise set to 7.5pt
                m1 = re.findall(r'\n[\w\.\- ]{14,}|[\w\.\- ]{14,}\n', run.text)
                m2 = re.findall(r'\w{9,}-[\w\.\- ]{14,}|\n(?:\w+-){4,}\w+|(?:\w+-){4,}\w+\n', run.text)
                m3 = re.findall(r'\n[\w\.\- ]{25,}|[\w\.\- ]{25,}\n', run.text)
                m4 = re.findall(r'\n[\w\.\- ]{36,}|[\w\.\- ]{36,}\n|(?:\w+ ){4,}\w+', run.text)
                m5 = re.findall(r'\n(?:\w{1,4}-){3}\w{1,3}|^(?:\w{1,4}-){3}\w{1,3}\n', run.text)
                if sum([len(m1), len(m2), len(m3), len(m4)]) - len(m5) > 6:
                    run.font.size = Pt(6)
                elif sum([len(m1), len(m2), len(m3), len(m4)]) - len(m5) > 3:
                    run.font.size = Pt(6.5)
                elif sum([len(m1), len(m2), len(m3), len(m4)]) > 2:
                    run.font.size = Pt(7)
                else:
                    run.font.size = Pt(7.5)
                cell.width = Inches(1)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell, top=0, start=14.4, bottom=0, end=14.4)
                cellsNum += 1
                # If this is not the last cell on the row, add a blank column/cell and set the width, then increment the column number and cell number
                if colNum != 12:
                    colNum += 1
                    cell = table.cell(rowNum, colNum)
                    cell.width = Inches(0.1)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    set_cell_margins(cell, top=0, start=14.4, bottom=0, end=14.4)
                    colNum += 1
                    cellsNum += 1
                # If this is the last column in the row, then we need to format the empty cells on the next row
                else:
                    colNum = 0
                    rowNum += 1
                    table.rows[rowNum].height = Inches(3)
                    table.rows[rowNum].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                    # Write the 13 empty cells for the row, alternating the widths, then move to next row
                    for n in range(0, 13):
                        cell = table.cell(rowNum, n)
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        set_cell_margins(cell, top=0, start=14.4, bottom=0, end=14.4)
                        if not n % 2:
                            cell.width = Inches(1)
                        else:
                            cell.width = Inches(0.1)
                        if cellsNum < 51:
                            cellsNum += 1
                    rowNum += 1

        # Set the margins for each section of the document
        sections = document.sections
        for section in sections:
            section.top_margin = Inches(0.81)
            section.bottom_margin = Inches(2)
            section.left_margin = Inches(0.45)
            section.right_margin = Inches(0.45)

        # Clean disallowed characters from the site name before trying to save the document
        illegal = '<>:"/\|?*'
        siteName = FME_MacroValues['SITENAME']
        for char in illegal:
            siteName = siteName.replace(char, '_')

        # Save the document
        document.save(f'{siteName}_Labels.docx')
        # Output the path to the created docx
        output = fmeobjects.FMEFeature()
        output.setGeometry(None)
        output.setAttribute('filecopy_source_dataset',
                            os.path.join(os.getcwd(), f'{FME_MacroValues["SITENAME"]}_Labels.docx'))
        output.setAttribute('filecopy_dest_dataset', os.getcwd())
        output.setAttribute('filecopy_dest_filename', f'{FME_MacroValues["SITENAME"]}_Labels.docx')
        self.pyoutput(output)